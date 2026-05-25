from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
import io
from app.models.report import Report
from app.core.logging import get_logger

logger = get_logger(__name__)

async def get_report(db: AsyncSession, report_id) -> Report | None:
    result = await db.execute(select(Report).where(Report.id == report_id))
    return result.scalar_one_or_none()

async def list_reports(db: AsyncSession, user_id, skip: int = 0, limit: int = 100) -> tuple[list[Report], int]:
    result = await db.execute(
        select(Report)
        .where(Report.user_id == user_id)
        .order_by(Report.created_at.desc())
        .offset(skip)
        .limit(limit)
    )
    reports = list(result.scalars().all())
    
    count_result = await db.execute(select(func.count(Report.id)).where(Report.user_id == user_id))
    total = count_result.scalar()
    
    return reports, total

from sqlalchemy import delete
from app.models.workflow import Workflow

async def delete_report(db: AsyncSession, report_id) -> bool:
    report = await get_report(db, report_id)
    if not report:
        return False
        
    workflow_id = report.workflow_id
    
    # First delete the report to avoid foreign key constraints
    result = await db.execute(delete(Report).where(Report.id == report_id))
    
    # Then cascade delete the workflow if it existed
    if workflow_id:
        from app.models.scraped_data import ScrapedData
        from app.models.agent_log import AgentLog
        from app.models.analytics import Analytics
        from app.models.embedding_metadata import EmbeddingMetadata
        
        await db.execute(delete(EmbeddingMetadata).where(EmbeddingMetadata.workflow_id == workflow_id))
        await db.execute(delete(Analytics).where(Analytics.workflow_id == workflow_id))
        await db.execute(delete(AgentLog).where(AgentLog.workflow_id == workflow_id))
        await db.execute(delete(ScrapedData).where(ScrapedData.workflow_id == workflow_id))
        await db.execute(delete(Workflow).where(Workflow.id == workflow_id))
        
    await db.commit()
    return result.rowcount > 0

def generate_pdf(report: Report) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, BaseDocTemplate, PageTemplate, Frame
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas
        import datetime
        
        buffer = io.BytesIO()
        
        # Define modern color palette
        COLOR_PRIMARY = colors.HexColor('#6366F1') # Indigo
        COLOR_SECONDARY = colors.HexColor('#0EA5E9') # Light Blue
        COLOR_DARK = colors.HexColor('#1E293B') # Slate 800
        COLOR_TEXT = colors.HexColor('#334155') # Slate 700
        COLOR_LIGHT = colors.HexColor('#F8FAFC') # Slate 50
        COLOR_ACCENT = colors.HexColor('#F43F5E') # Rose
        
        def header_footer(canvas_obj, doc):
            canvas_obj.saveState()
            # Header
            canvas_obj.setFillColor(COLOR_PRIMARY)
            canvas_obj.rect(0, letter[1] - 40, letter[0], 40, stroke=0, fill=1)
            canvas_obj.setFillColor(colors.white)
            canvas_obj.setFont("Helvetica-Bold", 12)
            canvas_obj.drawString(50, letter[1] - 25, "AgentFlow Market Intelligence")
            
            # Footer
            canvas_obj.setFillColor(COLOR_DARK)
            canvas_obj.rect(0, 0, letter[0], 30, stroke=0, fill=1)
            canvas_obj.setFillColor(colors.white)
            canvas_obj.setFont("Helvetica", 9)
            canvas_obj.drawString(50, 10, f"Generated {datetime.datetime.now().strftime('%Y-%m-%d')}")
            canvas_obj.drawRightString(letter[0] - 50, 10, f"Page {doc.page}")
            canvas_obj.restoreState()

        doc = BaseDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=60, bottomMargin=50)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='test', frames=frame, onPage=header_footer)
        doc.addPageTemplates([template])
        
        styles = getSampleStyleSheet()
        
        # Custom vibrant styles
        custom_title = ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            spaceBefore=100,
            spaceAfter=20,
            textColor=COLOR_PRIMARY,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=34
        )
        
        custom_subtitle = ParagraphStyle(
            name='CustomSubtitle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=100,
            textColor=COLOR_SECONDARY,
            alignment=TA_CENTER,
            fontName='Helvetica',
            leading=20
        )
        
        custom_heading = ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=24,
            spaceAfter=12,
            textColor=COLOR_DARK,
            fontName='Helvetica-Bold',
            borderPadding=6,
            borderWidth=0,
            bottomPadding=4
        )
        
        custom_normal = ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=COLOR_TEXT,
            fontName='Helvetica',
            leading=18
        )
        
        custom_bullet = ParagraphStyle(
            name='CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            leftIndent=25,
            textColor=COLOR_TEXT,
            fontName='Helvetica',
            leading=16,
            bulletIndent=10
        )
        
        flowables = []
        
        # --- COVER PAGE ---
        flowables.append(Spacer(1, 100))
        flowables.append(Paragraph("<b>INTELLIGENCE REPORT</b>", custom_subtitle))
        flowables.append(Paragraph(report.title or "Market Research Report", custom_title))
        
        flowables.append(Spacer(1, 50))
        flowables.append(Paragraph(f"<b>Prepared for:</b> Executive Team", ParagraphStyle('Centered', parent=custom_normal, alignment=TA_CENTER)))
        flowables.append(Paragraph(f"<b>Date:</b> {datetime.datetime.now().strftime('%B %d, %Y')}", ParagraphStyle('Centered', parent=custom_normal, alignment=TA_CENTER)))
        flowables.append(PageBreak())
        
        # --- EXECUTIVE SUMMARY ---
        flowables.append(Paragraph("Executive Summary", custom_heading))
        flowables.append(Paragraph(report.executive_summary or "No summary available.", custom_normal))
        
        # Load the dynamic advanced sections
        sections = {}
        try:
            if report.full_report:
                import json
                sections = json.loads(report.full_report)
        except Exception:
            pass
            
        # Dynamically render all sections
        for section_title, content in sections.items():
            flowables.append(Paragraph(section_title, custom_heading))
            
            if isinstance(content, list):
                for item in content:
                    flowables.append(Paragraph(f"<font color='#0EA5E9'>■</font> {item}", custom_bullet))
            elif isinstance(content, dict):
                for key, value in content.items():
                    if isinstance(value, list):
                        val_str = ', '.join(str(v) for v in value)
                        flowables.append(Paragraph(f"<b><font color='#6366F1'>{key.replace('_', ' ').title()}:</font></b> {val_str}", custom_normal))
                    else:
                        flowables.append(Paragraph(f"<b><font color='#6366F1'>{key.replace('_', ' ').title()}:</font></b> {value}", custom_normal))
            else:
                paragraphs = str(content).split('\n')
                for p in paragraphs:
                    if p.strip():
                        clean_p = p.strip().replace('**', '')
                        flowables.append(Paragraph(clean_p, custom_normal))

        if not sections and report.recommendations:
            flowables.append(Paragraph("Strategic Recommendations", custom_heading))
            for rec in report.recommendations:
                flowables.append(Paragraph(f"<font color='#0EA5E9'>■</font> {rec}", custom_bullet))
            
        doc.build(flowables)
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}")
        return b"%PDF-1.4 Error generating PDF"

def generate_docx(report: Report) -> bytes:
    try:
        import docx
        doc = docx.Document()
        
        doc.add_heading(report.title or "Market Research Report", 0)
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary or "No summary available.")
        
        doc.add_heading("Recommendations", level=1)
        for rec in (report.recommendations or []):
            doc.add_paragraph(rec, style='List Bullet')
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        logger.error(f"Failed to generate DOCX: {e}")
        return b"Error generating DOCX"
